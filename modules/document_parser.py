"""Extract text from files uploaded through the smart-import form."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Callable, List


class DocumentParseError(ValueError):
    """An upload is valid, but its contents cannot be read as text."""


def _load_pdf_reader() -> Callable:
    """Load the maintained PDF reader, with a compatibility fallback."""
    try:
        from pypdf import PdfReader

        return PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader

            return PdfReader
        except ImportError as exc:
            raise DocumentParseError(
                "当前环境缺少 PDF 解析依赖。请重新运行项目启动器，或安装 pypdf 后再试。"
            ) from exc


def _extract_pdf_text(file_bytes: bytes) -> str:
    if not file_bytes:
        raise DocumentParseError("上传的 PDF 文件为空。")

    reader_cls = _load_pdf_reader()
    try:
        try:
            reader = reader_cls(io.BytesIO(file_bytes), strict=False)
        except TypeError:
            # Older PyPDF2 releases do not accept ``strict``.
            reader = reader_cls(io.BytesIO(file_bytes))
    except Exception as exc:
        raise DocumentParseError(f"PDF 文件无法打开，可能已损坏或格式不受支持：{exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            decrypted = reader.decrypt("")
        except Exception as exc:
            raise DocumentParseError("PDF 受密码保护，暂时无法读取。请上传未加密版本。") from exc
        if decrypted == 0:
            raise DocumentParseError("PDF 受密码保护，暂时无法读取。请上传未加密版本。")

    page_text: List[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            # A malformed page should not prevent the remaining resume pages
            # from being imported.
            text = ""
        text = text.strip()
        if text:
            page_text.append(text)

    extracted = "\n\n".join(page_text).strip()
    if not extracted:
        raise DocumentParseError(
            "PDF 中没有提取到可读文字。若这是扫描件/图片 PDF，请先进行 OCR，"
            "或将文字复制到“粘贴文本”标签页。"
        )
    return extracted


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
    except ImportError as exc:
        raise DocumentParseError(
            "当前环境缺少 DOCX 解析依赖。请重新运行项目启动器后再试。"
        ) from exc
    except Exception as exc:
        raise DocumentParseError(f"DOCX 文件无法打开：{exc}") from exc

    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))

    extracted = "\n".join(paragraphs + table_text).strip()
    if not extracted:
        raise DocumentParseError(
            "DOCX 中没有可读取的文字内容，请尝试复制文本到“粘贴文本”标签页。"
        )
    return extracted


def _extract_txt_text(file_bytes: bytes) -> str:
    if not file_bytes:
        raise DocumentParseError("上传的 TXT 文件为空。")
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            extracted = file_bytes.decode(encoding).strip()
            if extracted:
                return extracted
        except UnicodeDecodeError:
            continue
    extracted = file_bytes.decode("utf-8", errors="ignore").strip()
    if not extracted:
        raise DocumentParseError("TXT 文件中没有可读取的文字内容。")
    return extracted


def extract_document_text(filename: str, file_bytes: bytes) -> str:
    """Return readable text from a supported smart-import upload.

    ``file_bytes`` is used instead of the uploaded stream directly so the
    function remains deterministic across Streamlit reruns.
    """
    suffix = Path(filename or "").suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(file_bytes)
    if suffix == ".docx":
        return _extract_docx_text(file_bytes)
    if suffix == ".txt":
        return _extract_txt_text(file_bytes)
    raise DocumentParseError("暂不支持该文件类型，请上传 PDF、DOCX 或 TXT 文件。")
